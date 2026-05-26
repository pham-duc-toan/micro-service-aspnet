"""Convert BAO_CAO_HTPT.md -> BAO_CAO_HTPT.docx with report-style formatting.

Pipeline:
1. pypandoc converts markdown -> docx (keeps headings, tables, code blocks).
2. python-docx post-processes the file: Times New Roman 13pt, line spacing 1.5,
   margins (top/bottom 2cm, left 3cm, right 2cm). Code blocks keep Consolas.
"""

from pathlib import Path

import pypandoc
from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

import os

ROOT = Path(__file__).resolve().parent
SRC = ROOT / "BAO_CAO_HTPT.md"
DST_PREFERRED = ROOT / "BAO_CAO_HTPT.docx"
DST_FALLBACK = ROOT / "BAO_CAO_HTPT_new.docx"


def pick_destination():
    """Use the preferred path if writable; otherwise fall back so a locked
    Word document does not block the build."""
    try:
        if DST_PREFERRED.exists():
            with open(DST_PREFERRED, "ab"):
                pass
        return DST_PREFERRED
    except PermissionError:
        return DST_FALLBACK


DST = pick_destination()

BODY_FONT = "Times New Roman"
CODE_FONT = "Consolas"
BODY_SIZE = Pt(13)


def convert_markdown():
    pypandoc.convert_file(
        str(SRC),
        "docx",
        outputfile=str(DST),
        extra_args=[
            "--from=gfm",
            "--toc",
            "--toc-depth=3",
            "--standalone",
        ],
    )


def set_run_font(run, font_name, size=None, bold=None):
    run.font.name = font_name
    if size is not None:
        run.font.size = size
    if bold is not None:
        run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        from docx.oxml import OxmlElement

        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), font_name)


def style_paragraph(par, is_code=False):
    fmt = par.paragraph_format
    if is_code:
        fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
    else:
        fmt.line_spacing = 1.5
        fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(6)
    for run in par.runs:
        if is_code:
            set_run_font(run, CODE_FONT, Pt(10))
        else:
            set_run_font(run, BODY_FONT, BODY_SIZE)


def style_document():
    doc = Document(str(DST))

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)

    for style_name in ("Normal", "Body Text", "List Paragraph"):
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue
        style.font.name = BODY_FONT
        style.font.size = BODY_SIZE
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            from docx.oxml import OxmlElement

            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rfonts.set(qn(attr), BODY_FONT)

    heading_sizes = {
        "Title": Pt(20),
        "Heading 1": Pt(16),
        "Heading 2": Pt(14),
        "Heading 3": Pt(13),
        "Heading 4": Pt(13),
    }
    for name, size in heading_sizes.items():
        try:
            style = doc.styles[name]
        except KeyError:
            continue
        style.font.name = BODY_FONT
        style.font.size = size
        style.font.bold = True
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            from docx.oxml import OxmlElement

            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rfonts.set(qn(attr), BODY_FONT)

    for par in doc.paragraphs:
        style_name = (par.style.name or "").lower()
        is_code = "source code" in style_name or "code" in style_name
        style_paragraph(par, is_code=is_code)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for par in cell.paragraphs:
                    style_paragraph(par)

    doc.save(str(DST))


def main():
    if not SRC.exists():
        raise SystemExit(f"Source not found: {SRC}")
    convert_markdown()
    style_document()
    print(f"OK -> {DST}")


if __name__ == "__main__":
    main()
